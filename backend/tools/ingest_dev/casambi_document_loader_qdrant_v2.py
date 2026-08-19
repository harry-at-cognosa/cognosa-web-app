###
# Document loader for Casambi Collection - Qdrant
# Multi-strategy chunking based on document classification
# Uses folder structure AND filename for classification
###

# Configuration
DOCUMENTS_PATH = "../../documents_casambi"  # Outside project hierarchy
QDRANT_URL = "localhost:6333"
COLLECTION_NAME = "casambi_collection"
TRANSFORMER_MODEL = "sentence-transformers/all-MiniLM-L6-v2"

# Chunking parameters per strategy (can be tuned)
CHUNK_PARAMS = {
    'cheat_sheet': {'chunk_size': 4000, 'overlap': 0},  # Single chunk
    'use_case': {'chunk_size': 800, 'overlap': 100},
    'manual': {'chunk_size': 1200, 'overlap': 200},
    'technical': {'chunk_size': 1000, 'overlap': 150},
    'specification': {'chunk_size': 600, 'overlap': 50},
}

# Default for unknown types
DEFAULT_CHUNK_SIZE = 1000
DEFAULT_CHUNK_OVERLAP = 150
MAX_PROCESSES = 16

import os
import re
from dataclasses import dataclass
from multiprocessing import Process, Queue, cpu_count
from threading import Thread
from pathlib import Path
from time import time
from typing import Literal, List, Dict, Optional, Tuple
from enum import Enum

import docx
from langchain_community.document_loaders import (
    PyPDFLoader,
    Docx2txtLoader,
    UnstructuredPowerPointLoader,
    UnstructuredHTMLLoader,
    TextLoader,
    CSVLoader,
    UnstructuredExcelLoader
)
from langchain_core.documents import Document
from langchain_text_splitters import RecursiveCharacterTextSplitter
from common import WORK_DIR, log
from common.parsed_url import ParsedUrl
from tasks_lib.vdb_lib.emb_models import HuggingFaceEmbeddings, EmbModels
from tasks_lib.vdb_lib.qdrant_ops import QdrantOps
from tasks_lib.vdb_lib.vdb_ops import VectorDBOps

DOCUMENTS_PATH = os.path.normpath(os.path.join(WORK_DIR, DOCUMENTS_PATH))

log.init(prefix='casambi_document_loader_qdrant')

# Document classification
class DocType(Enum):
    CHEAT_SHEET = "cheat_sheet"
    USE_CASE = "use_case"
    MANUAL = "manual"
    SYSTEM_OVERVIEW = "system_overview"
    SPECIFICATION = "specification"
    SECURITY_TECHNICAL = "security_technical"
    SUSTAINABILITY = "sustainability"
    PROCEDURAL_GUIDE = "procedural_guide"
    GENERAL_TECHNICAL = "general_technical"

@dataclass
class DocFile:
    path: str
    type: Literal['pdf', 'docx', 'pptx', 'html', 'txt', 'csv', 'xlsx']
    doc_category: Optional[str] = None  # Will be determined by classifier
    folder_category: Optional[str] = None  # Derived from folder structure

@dataclass
class ClassificationResult:
    doc_type: DocType
    confidence: float
    strategy: str
    classification_method: str  # 'folder', 'filename', or 'combined'


class CasambiDocumentClassifier:
    """Classify Casambi documents using folder structure and filename."""
    
    @staticmethod
    def classify(doc_path: str) -> ClassificationResult:
        """
        Classify document based on folder structure first, then filename.
        
        Priority:
        1. Folder structure (highest confidence)
        2. Folder + filename combination
        3. Filename only (fallback)
        """
        
        filename = os.path.basename(doc_path).lower()
        full_path = doc_path.lower()
        
        # Extract folder information
        path_parts = full_path.split(os.sep)
        
        # Method 1: Check folder structure (highest confidence)
        # Looking for folder names that indicate document type
        
        # CHEAT SHEETS - folder based
        if any(folder in full_path for folder in ['/cheat_sheet/', '/cheatsheet/', '/cheat-sheet/', 
                                                   '/quick_reference/', '/quick-reference/']):
            return ClassificationResult(
                doc_type=DocType.CHEAT_SHEET,
                confidence=0.95,
                strategy='single_chunk',
                classification_method='folder'
            )
        
        # USE CASES - folder based
        if any(folder in full_path for folder in ['/use_case/', '/usecase/', '/use-case/',
                                                   '/case_study/', '/case-study/', '/vertical/']):
            # Further refine by filename if in use_case folder
            confidence = 0.90
            if any(term in filename for term in ['hospitality', 'retail', 'office', 
                                                  'residential', 'outdoor', 'retrofit']):
                confidence = 0.95
            
            return ClassificationResult(
                doc_type=DocType.USE_CASE,
                confidence=confidence,
                strategy='semantic_section',
                classification_method='folder'
            )
        
        # MANUAL - folder based
        if any(folder in full_path for folder in ['/manual/', '/user_guide/', '/user-guide/',
                                                   '/documentation/']):
            # Check if it's the app manual specifically
            if 'app' in filename or 'user' in filename:
                return ClassificationResult(
                    doc_type=DocType.MANUAL,
                    confidence=0.95,
                    strategy='hierarchical',
                    classification_method='folder'
                )
            else:
                # Other manuals might be procedural guides
                return ClassificationResult(
                    doc_type=DocType.PROCEDURAL_GUIDE,
                    confidence=0.85,
                    strategy='procedural',
                    classification_method='folder'
                )
        
        # GENERAL TOPICS / TECHNICAL - folder based with filename refinement
        if any(folder in full_path for folder in ['/general_topic/', '/general-topic/', 
                                                   '/technical/', '/system/', '/spec/']):
            # Now use filename to determine specific type
            
            # Security documents
            if 'security' in filename:
                return ClassificationResult(
                    doc_type=DocType.SECURITY_TECHNICAL,
                    confidence=0.90,
                    strategy='technical_topic',
                    classification_method='combined'
                )
            
            # Sustainability documents
            if any(term in filename for term in ['sustainability', 'leed', 'well-being', 'wellbeing']):
                return ClassificationResult(
                    doc_type=DocType.SUSTAINABILITY,
                    confidence=0.90,
                    strategy='semantic_section',
                    classification_method='combined'
                )
            
            # System overview
            if 'overview' in filename or 'introduction' in filename:
                return ClassificationResult(
                    doc_type=DocType.SYSTEM_OVERVIEW,
                    confidence=0.85,
                    strategy='hierarchical',
                    classification_method='combined'
                )
            
            # Specification documents
            if any(term in filename for term in ['specification', 'spec', 'requirement']):
                return ClassificationResult(
                    doc_type=DocType.SPECIFICATION,
                    confidence=0.85,
                    strategy='specification',
                    classification_method='combined'
                )
            
            # iBeacon and other procedural guides
            if any(term in filename for term in ['ibeacon', 'setup', 'installation', 'configuration']):
                return ClassificationResult(
                    doc_type=DocType.PROCEDURAL_GUIDE,
                    confidence=0.85,
                    strategy='procedural',
                    classification_method='combined'
                )
            
            # Default for general topics folder
            return ClassificationResult(
                doc_type=DocType.GENERAL_TECHNICAL,
                confidence=0.70,
                strategy='standard',
                classification_method='folder'
            )
        
        # Method 2: Fallback to filename-only classification if no folder match
        # (Lower confidence since we're not using folder structure)
        
        # Cheat sheets - filename based
        if 'cheat' in filename or 'quick' in filename:
            return ClassificationResult(
                doc_type=DocType.CHEAT_SHEET,
                confidence=0.80,  # Lower confidence without folder
                strategy='single_chunk',
                classification_method='filename'
            )
        
        # Manual - filename based
        if 'manual' in filename and ('user' in filename or 'app' in filename):
            return ClassificationResult(
                doc_type=DocType.MANUAL,
                confidence=0.75,
                strategy='hierarchical',
                classification_method='filename'
            )
        
        # Use cases - filename based
        if any(term in filename for term in ['hospitality', 'retail', 'office', 'residential', 
                                             'outdoor', 'retrofit', 'refurbishment']):
            return ClassificationResult(
                doc_type=DocType.USE_CASE,
                confidence=0.75,
                strategy='semantic_section',
                classification_method='filename'
            )
        
        # Security - filename based
        if 'security' in filename:
            return ClassificationResult(
                doc_type=DocType.SECURITY_TECHNICAL,
                confidence=0.75,
                strategy='technical_topic',
                classification_method='filename'
            )
        
        # Sustainability - filename based
        if any(term in filename for term in ['sustainability', 'leed', 'well-being']):
            return ClassificationResult(
                doc_type=DocType.SUSTAINABILITY,
                confidence=0.75,
                strategy='semantic_section',
                classification_method='filename'
            )
        
        # iBeacon - filename based
        if 'ibeacon' in filename:
            return ClassificationResult(
                doc_type=DocType.PROCEDURAL_GUIDE,
                confidence=0.75,
                strategy='procedural',
                classification_method='filename'
            )
        
        # System overview - filename based
        if 'overview' in filename:
            return ClassificationResult(
                doc_type=DocType.SYSTEM_OVERVIEW,
                confidence=0.70,
                strategy='hierarchical',
                classification_method='filename'
            )
        
        # Specification - filename based
        if 'specification' in filename or 'spec' in filename:
            return ClassificationResult(
                doc_type=DocType.SPECIFICATION,
                confidence=0.70,
                strategy='specification',
                classification_method='filename'
            )
        
        # Default fallback
        return ClassificationResult(
            doc_type=DocType.GENERAL_TECHNICAL,
            confidence=0.50,
            strategy='standard',
            classification_method='default'
        )
    
    @staticmethod
    def extract_folder_category(doc_path: str) -> str:
        """Extract the category from folder structure."""
        path_lower = doc_path.lower()
        
        if '/cheat_sheet/' in path_lower or '/cheatsheet/' in path_lower:
            return 'cheat_sheets'
        elif '/use_case/' in path_lower or '/usecase/' in path_lower:
            return 'use_cases'
        elif '/manual/' in path_lower:
            return 'manual'
        elif '/general_topic/' in path_lower or '/technical/' in path_lower:
            return 'general_topics'
        else:
            # Try to extract from immediate parent folder
            parent_folder = os.path.basename(os.path.dirname(doc_path)).lower()
            if parent_folder in ['cheat_sheets', 'cheatsheets', 'use_cases', 'usecases', 
                               'manual', 'general_topics', 'technical']:
                return parent_folder.replace(' ', '_')
            return 'uncategorized'


class CasambiChunkingStrategies:
    """Different chunking strategies for Casambi documents."""
    
    @staticmethod
    def single_chunk_strategy(documents: List[Document]) -> List[Document]:
        """Keep entire document as single chunk (for cheat sheets)."""
        if not documents:
            return []
        
        # Combine all content into single document
        combined_content = "\n\n".join(doc.page_content for doc in documents)
        
        # Preserve original metadata from first doc and add strategy info
        metadata = documents[0].metadata.copy()
        metadata.update({
            'chunking_strategy': 'single_chunk',
            'total_pages': len(documents),
            'chunk_type': 'complete_document'
        })
        
        return [Document(page_content=combined_content, metadata=metadata)]
    
    @staticmethod
    def semantic_section_strategy(documents: List[Document], chunk_size: int = 800, 
                                 chunk_overlap: int = 100) -> List[Document]:
        """Chunk by semantic sections (for use cases, sustainability)."""
        chunks = []
        
        # Try to detect section boundaries
        for doc in documents:
            content = doc.page_content
            
            # Look for heading patterns
            sections = re.split(r'\n(?=[A-Z][A-Za-z\s]+:|\d+\.\s+[A-Z])', content)
            
            for i, section in enumerate(sections):
                if len(section.strip()) > 50:  # Minimum section size
                    metadata = doc.metadata.copy()
                    metadata.update({
                        'chunking_strategy': 'semantic_section',
                        'section_number': i + 1,
                        'section_type': CasambiChunkingStrategies._classify_section(section)
                    })
                    
                    # If section is too large, split it
                    if len(section) > chunk_size * 1.5:
                        splitter = RecursiveCharacterTextSplitter(
                            chunk_size=chunk_size,
                            chunk_overlap=chunk_overlap
                        )
                        sub_chunks = splitter.split_text(section)
                        for j, sub_chunk in enumerate(sub_chunks):
                            sub_metadata = metadata.copy()
                            sub_metadata['sub_chunk'] = j + 1
                            chunks.append(Document(page_content=sub_chunk, metadata=sub_metadata))
                    else:
                        chunks.append(Document(page_content=section, metadata=metadata))
        
        return chunks if chunks else documents
    
    @staticmethod
    def hierarchical_strategy(documents: List[Document], chunk_size: int = 1200,
                            chunk_overlap: int = 200) -> List[Document]:
        """Hierarchical chunking for manuals and overviews."""
        chunks = []
        
        # Build hierarchy from document structure
        current_h1 = None
        current_h2 = None
        current_content = []
        
        for doc in documents:
            content = doc.page_content
            lines = content.split('\n')
            
            for line in lines:
                # Detect heading levels (simplified - you might need to enhance this)
                if re.match(r'^#\s+', line) or (line.isupper() and len(line) > 5 and len(line) < 100):
                    # Major heading
                    if current_content:
                        chunks.append(CasambiChunkingStrategies._create_hierarchical_chunk(
                            current_h1, current_h2, current_content, doc.metadata
                        ))
                    current_h1 = line.strip('#').strip()
                    current_h2 = None
                    current_content = []
                    
                elif re.match(r'^##\s+', line) or (line.strip() and line[0].isupper() and 10 < len(line) < 80):
                    # Sub-heading
                    if current_content:
                        chunks.append(CasambiChunkingStrategies._create_hierarchical_chunk(
                            current_h1, current_h2, current_content, doc.metadata
                        ))
                    current_h2 = line.strip('#').strip()
                    current_content = []
                    
                else:
                    current_content.append(line)
            
            # Save last section
            if current_content:
                chunks.append(CasambiChunkingStrategies._create_hierarchical_chunk(
                    current_h1, current_h2, current_content, doc.metadata
                ))
        
        # If no structure detected, fall back to standard chunking
        if not chunks:
            splitter = RecursiveCharacterTextSplitter(
                chunk_size=chunk_size,
                chunk_overlap=chunk_overlap
            )
            return splitter.split_documents(documents)
        
        return chunks
    
    @staticmethod
    def technical_topic_strategy(documents: List[Document], chunk_size: int = 1000,
                                chunk_overlap: int = 150) -> List[Document]:
        """Topic-based chunking for technical documents."""
        chunks = []
        
        for doc in documents:
            content = doc.page_content
            
            # Look for technical topic markers
            topics = re.split(r'\n(?=(?:Protocol|Security|Authentication|Encryption|Architecture|Integration|API)\s*:)', 
                            content, flags=re.IGNORECASE)
            
            for i, topic in enumerate(topics):
                if len(topic.strip()) > 50:
                    # Extract topic name if present
                    topic_match = re.match(r'^([A-Za-z\s]+):\s*(.+)', topic, re.DOTALL)
                    if topic_match:
                        topic_name = topic_match.group(1).strip()
                        topic_content = topic_match.group(2).strip()
                    else:
                        topic_name = f"Topic_{i+1}"
                        topic_content = topic.strip()
                    
                    metadata = doc.metadata.copy()
                    metadata.update({
                        'chunking_strategy': 'technical_topic',
                        'topic': topic_name,
                        'topic_type': CasambiChunkingStrategies._classify_technical_topic(topic_name, topic_content)
                    })
                    
                    # Split if too large
                    if len(topic_content) > chunk_size * 1.5:
                        splitter = RecursiveCharacterTextSplitter(
                            chunk_size=chunk_size,
                            chunk_overlap=chunk_overlap
                        )
                        sub_chunks = splitter.split_text(topic_content)
                        for j, sub_chunk in enumerate(sub_chunks):
                            sub_metadata = metadata.copy()
                            sub_metadata['sub_chunk'] = j + 1
                            chunks.append(Document(page_content=sub_chunk, metadata=sub_metadata))
                    else:
                        chunks.append(Document(page_content=topic_content, metadata=metadata))
        
        return chunks if chunks else documents
    
    @staticmethod
    def procedural_strategy(documents: List[Document], chunk_size: int = 800,
                          chunk_overlap: int = 100) -> List[Document]:
        """Keep procedures together."""
        chunks = []
        
        for doc in documents:
            content = doc.page_content
            
            # Look for numbered steps or procedures
            procedures = re.split(r'\n(?=(?:Step|Procedure|Task|Configure|Setup)\s+\d+)', 
                                content, flags=re.IGNORECASE)
            
            for i, procedure in enumerate(procedures):
                if len(procedure.strip()) > 50:
                    # Check if this contains numbered steps
                    has_steps = bool(re.search(r'\n\s*\d+\.', procedure))
                    
                    metadata = doc.metadata.copy()
                    metadata.update({
                        'chunking_strategy': 'procedural',
                        'procedure_number': i + 1,
                        'has_numbered_steps': has_steps
                    })
                    
                    # Don't split procedures if possible
                    if len(procedure) > chunk_size * 2:
                        # Only split if really necessary
                        splitter = RecursiveCharacterTextSplitter(
                            chunk_size=chunk_size * 1.5,  # Larger chunks for procedures
                            chunk_overlap=chunk_overlap
                        )
                        sub_chunks = splitter.split_text(procedure)
                        for j, sub_chunk in enumerate(sub_chunks):
                            sub_metadata = metadata.copy()
                            sub_metadata['sub_chunk'] = j + 1
                            chunks.append(Document(page_content=sub_chunk, metadata=sub_metadata))
                    else:
                        chunks.append(Document(page_content=procedure, metadata=metadata))
        
        return chunks if chunks else documents
    
    @staticmethod
    def _create_hierarchical_chunk(h1: Optional[str], h2: Optional[str], 
                                  content: List[str], base_metadata: dict) -> Document:
        """Create a hierarchical chunk with metadata."""
        metadata = base_metadata.copy()
        metadata.update({
            'chunking_strategy': 'hierarchical',
            'section_h1': h1 or 'General',
            'section_h2': h2 or None,
            'hierarchy_level': 2 if h2 else 1
        })
        
        content_text = '\n'.join(content).strip()
        if h2:
            content_text = f"{h2}\n\n{content_text}"
        if h1 and not h2:
            content_text = f"{h1}\n\n{content_text}"
            
        return Document(page_content=content_text, metadata=metadata)
    
    @staticmethod
    def _classify_section(section_text: str) -> str:
        """Classify the type of section."""
        text_lower = section_text[:500].lower()
        
        if any(term in text_lower for term in ['benefit', 'advantage', 'improve']):
            return 'benefits'
        elif any(term in text_lower for term in ['challenge', 'problem', 'issue']):
            return 'challenges'
        elif any(term in text_lower for term in ['solution', 'solve', 'address']):
            return 'solutions'
        elif any(term in text_lower for term in ['case study', 'example', 'implementation']):
            return 'case_study'
        elif any(term in text_lower for term in ['introduction', 'overview', 'about']):
            return 'introduction'
        else:
            return 'general'
    
    @staticmethod
    def _classify_technical_topic(topic_name: str, content: str) -> str:
        """Classify technical topic type."""
        combined = f"{topic_name} {content[:200]}".lower()
        
        if any(term in combined for term in ['protocol', 'standard', 'specification']):
            return 'protocol'
        elif any(term in combined for term in ['security', 'encryption', 'authentication']):
            return 'security'
        elif any(term in combined for term in ['architecture', 'design', 'structure']):
            return 'architecture'
        elif any(term in combined for term in ['api', 'interface', 'integration']):
            return 'integration'
        else:
            return 'general_technical'


class LogThread(Thread):
    def __init__(self, total_files: int, result_queue: Queue):
        super().__init__()
        self.total_files = total_files
        self.result_queue = result_queue
        self.processed = 0
        self.total_chunks = 0

    def run(self):
        while True:
            msg = self.result_queue.get()
            if isinstance(msg, MsgExit):
                return
            if isinstance(msg, MsgFileFinished):
                self.processed += 1
                self.total_chunks += msg.chunks_created
                print(f"Files processed {self.processed}/{self.total_files}, Total chunks: {self.total_chunks}")


@dataclass
class MsgExit:
    exit: bool = True

@dataclass
class MsgFileFinished:
    path: str
    chunks_created: int = 0


class SingleDocumentProcess(Process):
    def __init__(self, process_index: int, msg_queue: Queue, result_queue: Queue):
        super().__init__()
        self.process_index = process_index
        self.msg_queue = msg_queue
        self.result_queue = result_queue
        self.collection_name = COLLECTION_NAME

    def load_document(self, doc_file: DocFile) -> List[Document]:
        """Load document using appropriate loader."""
        if doc_file.type == 'pdf':
            loader = PyPDFLoader(doc_file.path)
        elif doc_file.type == 'docx':
            loader = Docx2txtLoader(doc_file.path)
        elif doc_file.type == 'pptx':
            loader = UnstructuredPowerPointLoader(doc_file.path)
        elif doc_file.type == 'html':
            loader = UnstructuredHTMLLoader(doc_file.path)
        elif doc_file.type == 'txt':
            loader = TextLoader(doc_file.path, encoding='utf8', autodetect_encoding=True)
        elif doc_file.type == 'csv':
            loader = CSVLoader(doc_file.path, autodetect_encoding=True)
        elif doc_file.type == 'xlsx':
            loader = UnstructuredExcelLoader(doc_file.path)
        else:
            return []
        return loader.load()

    def apply_chunking_strategy(self, documents: List[Document], doc_file: DocFile) -> List[Document]:
        """Apply appropriate chunking strategy based on document classification."""
        
        # Classify document
        classification = CasambiDocumentClassifier.classify(doc_file.path)
        log.info(f"Classified as {classification.doc_type.value} with confidence {classification.confidence:.2f} (method: {classification.classification_method})")
        
        # Store classification in doc_file for later use
        doc_file.doc_category = classification.doc_type.value
        doc_file.folder_category = CasambiDocumentClassifier.extract_folder_category(doc_file.path)
        
        # Get chunking parameters
        if classification.doc_type.value in ['cheat_sheet']:
            return CasambiChunkingStrategies.single_chunk_strategy(documents)
        
        elif classification.doc_type.value in ['use_case', 'sustainability']:
            params = CHUNK_PARAMS.get('use_case', {})
            return CasambiChunkingStrategies.semantic_section_strategy(
                documents, 
                chunk_size=params.get('chunk_size', 800),
                chunk_overlap=params.get('overlap', 100)
            )
        
        elif classification.doc_type.value in ['manual', 'system_overview']:
            params = CHUNK_PARAMS.get('manual', {})
            return CasambiChunkingStrategies.hierarchical_strategy(
                documents,
                chunk_size=params.get('chunk_size', 1200),
                chunk_overlap=params.get('overlap', 200)
            )
        
        elif classification.doc_type.value in ['security_technical']:
            params = CHUNK_PARAMS.get('technical', {})
            return CasambiChunkingStrategies.technical_topic_strategy(
                documents,
                chunk_size=params.get('chunk_size', 1000),
                chunk_overlap=params.get('overlap', 150)
            )
        
        elif classification.doc_type.value in ['procedural_guide']:
            return CasambiChunkingStrategies.procedural_strategy(
                documents,
                chunk_size=800,
                chunk_overlap=100
            )
        
        elif classification.doc_type.value in ['specification']:
            params = CHUNK_PARAMS.get('specification', {})
            # Use standard splitter with smaller chunks for specifications
            splitter = RecursiveCharacterTextSplitter(
                chunk_size=params.get('chunk_size', 600),
                chunk_overlap=params.get('overlap', 50)
            )
            return splitter.split_documents(documents)
        
        else:
            # Default strategy
            splitter = RecursiveCharacterTextSplitter(
                chunk_size=DEFAULT_CHUNK_SIZE,
                chunk_overlap=DEFAULT_CHUNK_OVERLAP
            )
            return splitter.split_documents(documents)

    def enrich_metadata(self, chunks: List[Document], doc_file: DocFile, classification: ClassificationResult) -> List[Document]:
        """Add Casambi-specific metadata to chunks."""
        
        filename = os.path.basename(doc_file.path)
        file_path = doc_file.path
        
        # Extract folder path relative to documents root
        relative_path = os.path.relpath(doc_file.path, DOCUMENTS_PATH)
        folder_structure = os.path.dirname(relative_path)
        
        for i, chunk in enumerate(chunks):
            # Add common Casambi metadata
            chunk.metadata.update({
                'collection': 'casambi',
                'source_file': filename,
                'source_path': file_path,
                'relative_path': relative_path,
                'folder_structure': folder_structure,
                'folder_category': doc_file.folder_category,
                'doc_type': classification.doc_type.value,
                'doc_category': doc_file.doc_category,
                'classification_confidence': classification.confidence,
                'classification_method': classification.classification_method,
                'chunk_index': i,
                'total_chunks': len(chunks),
                'processing_date': time()
            })
        
        return chunks

    def process_file(self, doc_file: DocFile):
        """Process a single file with appropriate chunking strategy."""
        
        # Load document
        doc_list = self.load_document(doc_file)
        if not doc_list:
            log.error(f"Failed to load file: {doc_file.path}")
            return
        
        # Classify and apply chunking strategy
        classification = CasambiDocumentClassifier.classify(doc_file.path)
        split_docs = self.apply_chunking_strategy(doc_list, doc_file)
        
        if not split_docs:
            log.error(f"Failed to get chunks from file: {doc_file.path}")
            return
        
        # Enrich metadata
        split_docs = self.enrich_metadata(split_docs, doc_file, classification)
        
        log.debug(f"Split into {len(split_docs)} chunks using {classification.strategy} strategy")
        
        # Save to Qdrant
        self.qdrant_ops.save_to_qdrant(
            self.emb_obj.get_by_name(TRANSFORMER_MODEL), 
            self.collection_name, 
            split_docs
        )
        
        return len(split_docs)

    def run(self):
        log.init(prefix=f'casambi_loader_{self.process_index}')
        self.qdrant_ops = QdrantOps(ParsedUrl.from_url(QDRANT_URL))
        self.emb_obj = EmbModels()
        
        while True:
            msg: DocFile | MsgExit = self.msg_queue.get()
            if isinstance(msg, MsgExit):
                self.msg_queue.put(msg)
                return
            if isinstance(msg, DocFile):
                try:
                    chunks_created = self.process_file(msg)
                    self.result_queue.put(MsgFileFinished(msg.path, chunks_created or 0))
                except Exception as exc:
                    log.error(f"Exception for file {msg.path} : {exc}")
                    self.result_queue.put(MsgFileFinished(msg.path, 0))


class CasambiDocumentLoader:
    """Main loader class for Casambi collection."""
    
    def __init__(self):
        self.documents_folder = Path(DOCUMENTS_PATH)
        self.collection_name = COLLECTION_NAME
        self.vdb_ops = VectorDBOps('qdrant', QDRANT_URL)
        
        if error_msg := self.vdb_ops.check_url():
            print(f"Wrong {QDRANT_URL=} : {error_msg}")
            exit(-1)
            
        self.qdrant_ops = QdrantOps(self.vdb_ops.parsed_url)
        self.start_time = time()
        log.info(f"Initialized CasambiDocumentLoader with folder: {DOCUMENTS_PATH}")
        self.file_list: List[DocFile] = []
        self.process_list: List[SingleDocumentProcess] = []
        self.result_queue = Queue()

    def start(self):
        """Start the document loading process."""
        self.file_list = self.get_file_list()
        if not self.file_list:
            log.error("No files found to process")
            return
            
        log.info(f"Found {len(self.file_list)} Casambi documents to process")
        
        # Print document classification summary
        self.print_classification_summary()
        
        self.log_thread = LogThread(len(self.file_list), self.result_queue)
        self.log_thread.start()
        self.create_collection()
        self.start_processes()
        self.result_queue.put(MsgExit())
        self.log_thread.join()
        
        used_time = time() - self.start_time
        log.info(f"Casambi document processing completed! Took {used_time:.2f} seconds")
        log.info(f"Total chunks created: {self.log_thread.total_chunks}")

    def print_classification_summary(self):
        """Print summary of document classifications."""
        classifications = {}
        folder_structure = {}
        
        for doc_file in self.file_list:
            # Get classification
            result = CasambiDocumentClassifier.classify(doc_file.path)
            doc_type = result.doc_type.value
            
            # Get folder info
            relative_path = os.path.relpath(doc_file.path, DOCUMENTS_PATH)
            folder = os.path.dirname(relative_path) or 'root'
            
            # Track by doc type
            if doc_type not in classifications:
                classifications[doc_type] = []
            classifications[doc_type].append({
                'file': os.path.basename(doc_file.path),
                'folder': folder,
                'confidence': result.confidence,
                'method': result.classification_method
            })
            
            # Track by folder
            if folder not in folder_structure:
                folder_structure[folder] = []
            folder_structure[folder].append(os.path.basename(doc_file.path))
        
        print("\n" + "="*60)
        print("DOCUMENT CLASSIFICATION SUMMARY")
        print("="*60)
        
        print("\n📁 FOLDER STRUCTURE:")
        for folder, files in sorted(folder_structure.items()):
            print(f"\n{folder}/")
            for f in files[:5]:  # Show first 5
                print(f"  ├── {f}")
            if len(files) > 5:
                print(f"  └── ... and {len(files)-5} more files")
        
        print("\n📋 CLASSIFICATION BY TYPE:")
        for doc_type, docs in sorted(classifications.items()):
            print(f"\n{doc_type}: {len(docs)} documents")
            # Group by classification method
            by_method = {}
            for doc in docs:
                method = doc['method']
                if method not in by_method:
                    by_method[method] = []
                by_method[method].append(doc)
            
            for method, method_docs in by_method.items():
                print(f"  [{method}]: {len(method_docs)} docs")
                for doc in method_docs[:2]:  # Show first 2
                    print(f"    - {doc['file'][:50]} (conf: {doc['confidence']:.2f})")
        
        print("\n" + "="*60 + "\n")

    def create_collection(self):
        """Create Qdrant collection if it doesn't exist."""
        if self.qdrant_ops.collection_exists(self.collection_name):
            log.info(f"Collection {self.collection_name} already exists")
            
            # Optionally delete and recreate for clean slate
            response = input("Collection exists. Delete and recreate? (y/n): ")
            if response.lower() == 'y':
                self.qdrant_ops.delete_collection(self.collection_name)
                log.info(f"Deleted existing collection {self.collection_name}")
            else:
                return
                
        log.info(f"Creating collection {self.collection_name}...")
        emb_obj = HuggingFaceEmbeddings(model_name=TRANSFORMER_MODEL)
        self.qdrant_ops.create_collection(emb_obj, self.collection_name)

    def start_processes(self):
        """Start multiprocessing for document loading."""
        # For Casambi, we might want fewer processes since chunking is more complex
        process_count = max(min(cpu_count() - 1, len(self.file_list)), 1)
        process_count = min(process_count, min(int(MAX_PROCESSES), 8))  # Cap at 8 for Casambi
        
        log.info(f"Starting {process_count} loader processes...")
        msg_queue = Queue()
        
        for process_index in range(1, process_count + 1):
            self.process_list.append(SingleDocumentProcess(process_index, msg_queue, self.result_queue))
            self.process_list[-1].start()
            
        for doc_file in self.file_list:
            msg_queue.put(doc_file)

        msg_queue.put(MsgExit())
        for document_process in self.process_list:
            document_process.join()

    def get_file_list(self) -> List[DocFile]:
        """Get list of Casambi documents to process."""
        file_list: List[DocFile] = []
        
        # Walk through documents folder
        for root, _, files in os.walk(self.documents_folder):
            for f_name in files:
                f_path = os.path.normpath(os.path.join(root, f_name))
                f_ext = f_path.split('.')[-1].lower()
                
                # Skip non-document files
                if f_ext == 'pdf':
                    doc_file = DocFile(f_path, 'pdf')
                elif f_ext == 'docx':
                    doc_file = DocFile(f_path, 'docx')
                elif f_ext == 'pptx':
                    doc_file = DocFile(f_path, 'pptx')
                elif f_ext in ('htm', 'html'):
                    doc_file = DocFile(f_path, 'html')
                elif f_ext in ('txt', 'md'):
                    doc_file = DocFile(f_path, 'txt')
                elif f_ext == 'csv':
                    doc_file = DocFile(f_path, 'csv')
                elif f_ext == 'xlsx':
                    doc_file = DocFile(f_path, 'xlsx')
                else:
                    continue
                
                # Extract folder category
                doc_file.folder_category = CasambiDocumentClassifier.extract_folder_category(f_path)
                file_list.append(doc_file)
        
        # Sort for consistent processing
        file_list.sort(key=lambda x: x.path)
        
        return file_list


if __name__ == "__main__":
    loader = CasambiDocumentLoader()
    loader.start()
